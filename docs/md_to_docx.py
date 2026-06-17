from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_shading(cell, fill_color):
    """Set cell background color using hex string like '1A477A'"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2d, 0x5f, 0x8a)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x3d, 0x7a, 0x9a)
            run.font.bold = True
        elif level == 4:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4d, 0x8a, 0xaa)
            run.font.bold = True
    return p

def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
    doc.add_paragraph()

def add_table_from_markdown(doc, lines):
    header = lines[0].strip()
    data_rows = [l.strip() for l in lines[2:]]
    headers = [c.strip() for c in header.split('|') if c.strip()]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1A477A')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = 'Arial'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(data_rows):
        cells = [c.strip() for c in row.split('|') if c.strip()]
        for c_idx, val in enumerate(cells):
            if c_idx < len(headers):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'

    doc.add_paragraph()

# Read markdown
with open('docs/API_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Title page
title = doc.add_paragraph()
title_run = title.add_run('RSDWAY API Documentation')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)
title_run.font.name = 'Arial'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

title2 = doc.add_paragraph()
title2_run = title2.add_run('Arabic/English Bilingual Developer Guide\nfor SFDA DTTS (Rasid) Integration')
title2_run.font.size = Pt(14)
title2_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
title2_run.font.name = 'Arial'
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('Base URL: https://app.rsdway.com/api')
info_run.font.size = Pt(11)
info_run.font.color.rgb = RGBColor(0x2d, 0x5f, 0x8a)
info_run.font.name = 'Courier New'

doc.add_page_break()

# Parse markdown
lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []
in_table = False
table_lines = []

while i < len(lines):
    line = lines[i]

    if i == 0 and line.startswith('# RSDWAY API Documentation'):
        i += 1
        continue

    if line.startswith('# Table of Contents'):
        while i < len(lines) and not lines[i].strip().startswith('---'):
            i += 1
        i += 1
        continue

    if line.startswith('```'):
        if in_code_block:
            add_code_block(doc, code_lines)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    if '|' in line and not line.startswith('#') and not line.startswith('>'):
        if not in_table:
            in_table = True
            table_lines = []
        table_lines.append(line)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_markdown(doc, table_lines)
            table_lines = []
            in_table = False

    if line.strip() == '---':
        i += 1
        continue

    if line.startswith('# '):
        text = line[2:].strip()
        add_heading_custom(doc, text, level=1)
        i += 1
        continue
    elif line.startswith('## '):
        text = line[3:].strip()
        add_heading_custom(doc, text, level=2)
        i += 1
        continue
    elif line.startswith('### '):
        text = line[4:].strip()
        add_heading_custom(doc, text, level=3)
        i += 1
        continue
    elif line.startswith('#### '):
        text = line[5:].strip()
        add_heading_custom(doc, text, level=4)
        i += 1
        continue

    if line.startswith('> '):
        text = line[2:].strip()
        text = re.sub(r'\*\*', '', text)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        i += 1
        continue

    if line.strip():
        # Process inline formatting
        text = line.strip()
        p = doc.add_paragraph()

        # Split by bold and inline code
        parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xe8, 0x3e, 0x3e)
            else:
                run = p.add_run(part)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    else:
        # Empty line - skip to keep spacing clean
        pass

    i += 1

# Save
doc.save('docs/RSDWAY_API_Documentation.docx')
print('Document saved successfully: docs/RSDWAY_API_Documentation.docx')
